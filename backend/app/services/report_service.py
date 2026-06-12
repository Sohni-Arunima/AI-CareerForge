import os

from docx import Document




def generate_career_report(

        workspace_id,

        workspace

):


    folder = "career_reports"



    if not os.path.exists(folder):


        os.makedirs(folder)




    path = (

        folder

        +

        "/career_report_"

        +

        workspace_id

        +

        ".docx"

    )




    doc = Document()



    doc.add_heading(

        "AI CareerForge Intelligence Report",

        level=0

    )




    doc.add_heading(

        "Candidate Resume Summary",

        level=1

    )



    doc.add_paragraph(

        workspace.get(

            "resume_text",

            ""

        )[:1000]

    )






    results = workspace.get(

        "results",

        {}

    )






    for feature,result in results.items():



        doc.add_heading(

            feature.replace(

                "_",

                " "

            ).title(),

            level=1

        )




        doc.add_paragraph(

            str(result)

        )







    doc.add_heading(

        "Final Recommendation",

        level=1

    )




    doc.add_paragraph(

        "Continue improving missing skills, strengthen projects, and optimize profile based on analysis."

    )





    doc.save(path)





    return {


        "report_file":

        path,


        "message":

        "Career Intelligence Report generated successfully"

    }