from docx import Document

from docx.shared import Pt, RGBColor

import os
from docx2pdf import convert



def create_resume_docx(
        data,
        template
):


    doc = Document()


    style = doc.styles["Normal"]

    style.font.name = "Calibri"

    style.font.size = Pt(11)



    # -----------------
    # TEMPLATE COLORS
    # -----------------

    colors = {

        "professional": RGBColor(30, 64, 175),

        "tech": RGBColor(37, 99, 235),

        "fresher": RGBColor(14, 165, 233),

        "executive": RGBColor(120, 53, 15),

        "data": RGBColor(8, 145, 178),

        "research": RGBColor(88, 28, 135),

        "creative": RGBColor(219, 39, 119),

        "ats": RGBColor(0, 0, 0)

    }



    theme = colors.get(
        template,
        colors["professional"]
    )



    # NAME HEADER

    title = doc.add_heading(
        data.get("name","Your Name"),
        level=0
    )


    title.runs[0].font.color.rgb = theme




    doc.add_paragraph(
        data.get("role","")
    )



    sections = [

        "summary",

        "skills",

        "experience",

        "projects",

        "education",

        "certifications"

    ]



    for section in sections:


        if section in data:


            heading = doc.add_heading(
                section.upper(),
                level=2
            )


            heading.runs[0].font.color.rgb = theme


            doc.add_paragraph(
                data[section]
            )



    if not os.path.exists(
        "generated_resumes"
    ):

        os.makedirs(
            "generated_resumes"
        )



        path = (

        "generated_resumes/"
        +
        template
        +
        "_resume.docx"

    )


    doc.save(path)



    pdf_path = path.replace(

        ".docx",

        ".pdf"

    )



    try:


        convert(

            path,

            pdf_path

        )


    except Exception as e:


        pdf_path = "PDF conversion failed"




    return {

        "docx": path,

        "pdf": pdf_path

    }